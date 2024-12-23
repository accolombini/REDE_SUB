'''
    REDE_SUB.docx

        ||> Objetivo: suporte à coleta e análise de dados de campo

            |> Gera automaticamente um arquivo CSV de template para os dados de campo (dados_campo.csv), caso não exista.
            |> Processa os dados coletados no campo e gera um relatório preliminar em PDF (relatorio_preliminar_inspecao.pdf).
'''

# Import de bibliotecas

import os
import pandas as pd
import pdfplumber
from docx import Document
from fpdf import FPDF
import re

# Funções para análise de documentos


def extract_text_from_pdf(pdf_path):
    """
    Extrai texto de um arquivo PDF.

    Args:
        pdf_path (str): Caminho para o arquivo PDF.

    Returns:
        str: Texto extraído do PDF.
    """
    if not os.path.exists(pdf_path):
        print("Arquivo PDF não encontrado. Gerando template...")
        create_pdf_template(pdf_path)

    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    """
    Extrai texto de um arquivo DOCX.

    Args:
        docx_path (str): Caminho para o arquivo DOCX.

    Returns:
        str: Texto extraído do DOCX.
    """
    if not os.path.exists(docx_path):
        print("Arquivo DOCX não encontrado. Gerando template...")
        create_docx_template(docx_path)

    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def create_docx_template(output_path):
    """
    Gera um arquivo DOCX template para simular dados de melhores práticas.

    Args:
        output_path (str): Caminho para salvar o arquivo DOCX.
    """
    doc = Document()
    doc.add_heading(
        "Melhores Práticas para Inspeção de Redes Subterrâneas", level=1)
    doc.add_paragraph("1. Inspeção visual detalhada das redes subterrâneas.")
    doc.add_paragraph("2. Medições periódicas com instrumentos calibrados.")
    doc.add_paragraph(
        "3. Registro das condições dos ativos e manutenção preventiva.")
    doc.add_paragraph(
        "4. Uso de tecnologias avançadas para detecção de falhas.")
    doc.save(output_path)
    print(f"Template DOCX gerado em: {output_path}")


def create_pdf_template(output_path):
    """
    Gera um arquivo PDF template para simular dados de melhores práticas.

    Args:
        output_path (str): Caminho para salvar o arquivo PDF.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Melhores Práticas para Inspeção de Redes Subterrâneas",
             ln=True, align="C")
    pdf.ln(10)
    pdf.cell(
        0, 10, txt="1. Inspeção visual detalhada das redes subterrâneas.", ln=True)
    pdf.cell(0, 10, txt="2. Medições periódicas com instrumentos calibrados.", ln=True)
    pdf.cell(
        0, 10, txt="3. Registro das condições dos ativos e manutenção preventiva.", ln=True)
    pdf.cell(
        0, 10, txt="4. Uso de tecnologias avançadas para detecção de falhas.", ln=True)
    pdf.output(output_path)
    print(f"Template PDF gerado em: {output_path}")


def create_field_data_template(output_path):
    """
    Gera um arquivo CSV template para simular dados de campo coletados.

    Args:
        output_path (str): Caminho para salvar o arquivo CSV.
    """
    data = {
        "ID_Ativo": ["A001", "A002", "A003"],
        "Data_Inspecao": ["2024-12-20", "2024-12-21", "2024-12-22"],
        "Condicao": ["Bom", "Regular", "Ruim"],
        "Medicoes": ["10.5", "8.3", "15.2"],
        "Observacoes": ["Sem anomalias.", "Pequenas rachaduras.", "Substituição recomendada."]
    }
    df = pd.DataFrame(data)
    df.to_csv(output_path, index=False)
    print(f"Template de dados de campo gerado em: {output_path}")


def analyze_text_with_keywords(text, keywords):
    """
    Analisa o texto para identificar palavras-chave relacionadas a inspeções.

    Args:
        text (str): Texto a ser analisado.
        keywords (list): Lista de palavras-chave para buscar no texto.

    Returns:
        dict: Dicionário com contagem de palavras-chave encontradas.
    """
    word_counts = {key: len(re.findall(f"\\b{key}\\b", text.lower()))
                   for key in keywords}
    return word_counts


def generate_template(output_path, data):
    """
    Gera um template automatizado de checklist ou protocolo em PDF.

    Args:
        output_path (str): Caminho para salvar o PDF.
        data (dict): Dados a serem incluídos no template.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", size=16, style="B")
    pdf.cell(200, 10, txt="Protocolo de Inspeção - Light", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("Arial", size=12)
    for key, value in data.items():
        pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)
        pdf.ln(5)

    pdf.output(output_path)
    print(f"Template de checklist gerado em: {output_path}")


def process_field_data(field_data_path):
    """
    Processa dados de campo coletados para gerar relatórios preliminares.

    Args:
        field_data_path (str): Caminho para o arquivo CSV com dados de campo.

    Returns:
        pd.DataFrame: Dados processados para análise e relatórios.
    """
    if not os.path.exists(field_data_path):
        print("Arquivo de dados de campo não encontrado. Gerando template...")
        create_field_data_template(field_data_path)

    return pd.read_csv(field_data_path)


def main():
    """
    Executa o script principal para auxiliar na coleta e análise de dados de campo.
    """
    base_path = "/Users/accol/Library/Mobile Documents/com~apple~CloudDocs/UNIVERSIDADES/UFF/PROJETOS/LIGHT/REDE_ATIVOS/REDE_SUB/script/DADOS"

    # Caminho para os dados de campo
    field_data_path = os.path.join(base_path, "dados_campo.csv")

    print("Processando dados de campo...")
    field_data = process_field_data(field_data_path)

    # Gerar relatório preliminar
    output_pdf_path = os.path.join(
        base_path, "relatorio_preliminar_inspecao.pdf")
    generate_template(output_pdf_path, {
        "Dados de Campo Coletados": field_data.to_dict(orient="records")
    })


if __name__ == "__main__":
    main()

'''
    REDE-SUB.py

    ||> Objetivo: apoio indireto, auxiliando no processamento e análise de informações para fundamentar o desenvolvimento do protocolo de inspeção.

    ||> Pesquisar Melhores Práticas
       
        |> Extrair informações de documentos ou relatórios existentes em formato PDF, DOCX ou CSV, usando bibliotecas como PyPDF2, pdfplumber, ou python-docx.
        Analisar texto: Identificar padrões ou palavras-chave relacionadas a inspeções de redes subterrâneas utilizando NLTK ou spaCy.

    ||> Propor Protocolo Detalhado
        Gerar templates automatizados para o protocolo, onde informações específicas podem ser preenchidas dinamicamente.
        Criar checklists automatizadas em formato PDF ou Excel para serem usadas no campo.
    
    ||> Adaptar o Protocolo às Condições da Light
        Processar dados históricos da Light, como:
            Falhas passadas: Identificar padrões ou locais mais críticos para inspeção.
            Inventário de Instrumentos: Automatizar a análise de disponibilidade de recursos e equipamentos.
'''

# Import bibliotecas

import os
import pandas as pd
import pdfplumber
from docx import Document
from fpdf import FPDF
import re

# Funções para análise de documentos


def extract_text_from_pdf(pdf_path):
    """
    Extrai texto de um arquivo PDF.

    Args:
        pdf_path (str): Caminho para o arquivo PDF.

    Returns:
        str: Texto extraído do PDF.
    """
    if not os.path.exists(pdf_path):
        print("Arquivo PDF não encontrado. Gerando template...")
        create_pdf_template(pdf_path)

    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    """
    Extrai texto de um arquivo DOCX.

    Args:
        docx_path (str): Caminho para o arquivo DOCX.

    Returns:
        str: Texto extraído do DOCX.
    """
    if not os.path.exists(docx_path):
        print("Arquivo DOCX não encontrado. Gerando template...")
        create_docx_template(docx_path)

    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def create_docx_template(output_path):
    """
    Gera um arquivo DOCX template para simular dados de melhores práticas.

    Args:
        output_path (str): Caminho para salvar o arquivo DOCX.
    """
    doc = Document()
    doc.add_heading(
        "Melhores Práticas para Inspeção de Redes Subterrâneas", level=1)
    doc.add_paragraph("1. Inspeção visual detalhada das redes subterrâneas.")
    doc.add_paragraph("2. Medições periódicas com instrumentos calibrados.")
    doc.add_paragraph(
        "3. Registro das condições dos ativos e manutenção preventiva.")
    doc.add_paragraph(
        "4. Uso de tecnologias avançadas para detecção de falhas.")
    doc.save(output_path)
    print(f"Template DOCX gerado em: {output_path}")


def create_pdf_template(output_path):
    """
    Gera um arquivo PDF template para simular dados de melhores práticas.

    Args:
        output_path (str): Caminho para salvar o arquivo PDF.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Melhores Práticas para Inspeção de Redes Subterrâneas",
             ln=True, align="C")
    pdf.ln(10)
    pdf.cell(
        0, 10, txt="1. Inspeção visual detalhada das redes subterrâneas.", ln=True)
    pdf.cell(0, 10, txt="2. Medições periódicas com instrumentos calibrados.", ln=True)
    pdf.cell(
        0, 10, txt="3. Registro das condições dos ativos e manutenção preventiva.", ln=True)
    pdf.cell(
        0, 10, txt="4. Uso de tecnologias avançadas para detecção de falhas.", ln=True)
    pdf.output(output_path)
    print(f"Template PDF gerado em: {output_path}")


def create_historical_data_template(output_path):
    """
    Gera um arquivo CSV template para simular dados históricos.

    Args:
        output_path (str): Caminho para salvar o arquivo CSV.
    """
    data = {
        "Tipo_Falha": ["Curto-circuito", "Defeito de isolação", "Falha mecânica"],
        "Local_Falha": ["Centro", "Zona Norte", "Zona Sul"],
        "Instrumento": ["Megômetro", "Multímetro", "Detector de falhas"]
    }
    df = pd.DataFrame(data)
    df.to_csv(output_path, index=False)
    print(f"Template de dados históricos gerado em: {output_path}")


def analyze_text_with_keywords(text, keywords):
    """
    Analisa o texto para identificar palavras-chave relacionadas a inspeções.

    Args:
        text (str): Texto a ser analisado.
        keywords (list): Lista de palavras-chave para buscar no texto.

    Returns:
        dict: Dicionário com contagem de palavras-chave encontradas.
    """
    word_counts = {key: len(re.findall(f"\\b{key}\\b", text.lower()))
                   for key in keywords}
    return word_counts


def generate_template(output_path, data):
    """
    Gera um template automatizado de checklist ou protocolo em PDF.

    Args:
        output_path (str): Caminho para salvar o PDF.
        data (dict): Dados a serem incluídos no template.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", size=16, style="B")
    pdf.cell(200, 10, txt="Protocolo de Inspeção - Light", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("Arial", size=12)
    for key, value in data.items():
        pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)
        pdf.ln(5)

    pdf.output(output_path)
    print(f"Template de checklist gerado em: {output_path}")


def process_historical_data(historical_data_path):
    """
    Processa dados históricos para identificar padrões de falhas e inventário de instrumentos.

    Args:
        historical_data_path (str): Caminho para o arquivo CSV com dados históricos.

    Returns:
        dict: Resumo dos padrões e inventário analisados.
    """
    if not os.path.exists(historical_data_path):
        print("Arquivo de dados históricos não encontrado. Gerando template...")
        create_historical_data_template(historical_data_path)

    df = pd.read_csv(historical_data_path)
    patterns = {
        "Falhas mais comuns": df["Tipo_Falha"].value_counts().to_dict(),
        "Locais mais críticos": df["Local_Falha"].value_counts().to_dict(),
        "Instrumentos necessários": df["Instrumento"].value_counts().to_dict()
    }
    return patterns


def main():
    """
    Executa o script principal para auxiliar no desenvolvimento do protocolo de inspeção.
    """
    base_path = "/Users/accol/Library/Mobile Documents/com~apple~CloudDocs/UNIVERSIDADES/UFF/PROJETOS/LIGHT/REDE_ATIVOS/REDE_SUB/script/DADOS"

    # Extrair texto de relatórios
    pdf_path = os.path.join(base_path, "melhores_praticas.pdf")
    docx_path = os.path.join(base_path, "melhores_praticas.docx")

    print("Extraindo texto de relatórios...")
    pdf_text = extract_text_from_pdf(pdf_path)
    docx_text = extract_text_from_docx(docx_path)

    # Analisar texto
    keywords = ["inspeção", "medição", "instrumento", "rede", "subterrânea"]
    print("Analisando texto para palavras-chave...")
    keyword_counts_pdf = analyze_text_with_keywords(pdf_text, keywords)
    keyword_counts_docx = analyze_text_with_keywords(docx_text, keywords)

    # Processar dados históricos
    historical_data_path = os.path.join(base_path, "dados_historicos.csv")
    print("Processando dados históricos...")
    historical_patterns = process_historical_data(historical_data_path)

    # Gerar checklist ou protocolo
    output_pdf_path = os.path.join(base_path, "protocolo_inspecao.pdf")
    checklist_data = {
        "Palavras-chave PDF": keyword_counts_pdf,
        "Palavras-chave DOCX": keyword_counts_docx,
        "Falhas mais comuns": historical_patterns["Falhas mais comuns"],
        "Locais mais críticos": historical_patterns["Locais mais críticos"],
        "Instrumentos necessários": historical_patterns["Instrumentos necessários"]
    }
    generate_template(output_pdf_path, checklist_data)


if __name__ == "__main__":
    main()
